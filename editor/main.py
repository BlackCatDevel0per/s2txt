from PyQt5.QtWidgets import QFileDialog, QMessageBox
from actions import UIC

import os

import sys

from docx import Document

from pathlib import Path


class Editor(UIC):
    def __init__(self, parent=None):
        super().__init__()

        self.current_file = None

    # Text Editor Code
    #################################################################################################

    def save_current_text(self):
        try:
            text_from_editor = self.textEdit.toPlainText()
            if text_from_editor is not "" and text_from_editor is not None and self.current_file is not None and os.path.isfile(self.current_file):
                if Path(self.current_file).suffix == ".txt":
                    with open(self.current_file, "w") as file:
                        text = self.textEdit.toPlainText()
                        file.write(text)
                elif Path(self.current_file).suffix == ".docx":
                    document = Document()
                    text = self.textEdit.toPlainText()
                    document.add_paragraph(text)
                    document.save(self.current_file)
            elif self.current_file is "" or self.current_file is None:
                # self.InformationUi("Файл не открыт!")
                SaveDialogData = QFileDialog.getSaveFileName(
                    self, 'Сохранить файл', os.path.expanduser("~"), "*.txt")
                if os.name == "nt":
                    SaveDialogData = SaveDialogData[0]
                elif os.name == "posix": # Fix file saving on linux
                    SaveDialogData = SaveDialogData[0] + SaveDialogData[1][1::]
                # print(SaveDialogData)

                if SaveDialogData is not "" and not os.path.isfile(SaveDialogData):
                    with open(SaveDialogData, "w") as file:
                        text = self.textEdit.toPlainText()
                        file.write(text)
                        #*#
                        self.current_file = SaveDialogData
                        self.action_close.setEnabled(True)
                        self.setWindowTitle(self.app_title + f" {self.current_file}")
                elif os.path.isfile(SaveDialogData):
                    self.WarningUi("Такой файл существует!")
        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))

    def file_close(self):
        try:
            if Path(self.current_file).suffix != ".docx":
                with open(self.current_file, "r") as file:
                    if file.read() == self.textEdit.toPlainText():
                        self.current_file = None
                        self.textEdit.clear()
                        self.setWindowTitle(self.app_title)
                        self.action_close.setEnabled(False)
                    elif file.read() != self.textEdit.toPlainText():
                        self.QuestionUi("Файл не сохранён!\nВыйти?")
                        if self.qui == QMessageBox.Ok:
                            self.current_file = None
                            self.textEdit.clear()
                            self.setWindowTitle(self.app_title)
                            self.action_close.setEnabled(False)
                    ###

            elif Path(self.current_file).suffix == ".docx":
                self.QuestionUi(
                    "Вы точно хотите закрыть файл? \n(проверьте изменения файла docx)")
                if self.qui == QMessageBox.Ok:
                    self.current_file = None
                    self.textEdit.clear()
                    self.setWindowTitle(self.app_title)
                    self.action_close.setEnabled(False)

        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))

    def quit_accept(self):
        try:
            fsuffix = Path(self.current_file).suffix
        except TypeError:
            fsuffix = ''

        try:
            if fsuffix == ".txt":
                if self.current_file is "" or self.current_file is None and self.textEdit.toPlainText() is "" or self.textEdit.toPlainText() is None:
                    sys.exit()
                    # self.aboutToQuit.connect(lambda: sys.exit())
                elif self.textEdit.toPlainText() is not "" and self.textEdit.toPlainText() is not None:
                    self.QuestionUi(
                        "Есть несохранённые изменения! \nВы точно хотите выйти?")
                    if self.qui == QMessageBox.Ok:
                        sys.exit()
            elif fsuffix == ".docx":
                self.QuestionUi("Вы точно хотите выйти?")
                if self.qui == QMessageBox.Ok:
                    sys.exit()
            elif fsuffix == '':
                self.QuestionUi("Вы точно хотите выйти?")
                if self.qui == QMessageBox.Ok:
                    sys.exit()

        except Exception as err:
            self.CriticalUi(str(err))

    ##############

    def save_text_as(self):
        try:
            SaveDialogData = QFileDialog.getSaveFileName(
                self, 'Сохранить файл', os.path.expanduser("~"), "*.txt")
            if os.name == "nt":
                SaveDialogData = SaveDialogData[0]
            elif os.name == "posix": # Fix file saving on linux
                SaveDialogData = SaveDialogData[0] + SaveDialogData[1][1::]
            # print(SaveDialogData)

            if SaveDialogData is not "" and not os.path.isfile(SaveDialogData):
                with open(SaveDialogData, "w") as file:
                    text = self.textEdit.toPlainText()
                    file.write(text)
            elif os.path.isfile(SaveDialogData):
                self.WarningUi("Такой файл существует!")

        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))

    def open_text(self):
        try:
            OpenDialogData = QFileDialog.getOpenFileName(
                self, 'Открыть файл', os.path.expanduser("~"), "*.txt")
            OpenDialogData = OpenDialogData[0]
            # print(OpenDialogData)

            if OpenDialogData is not "":
                with open(OpenDialogData, "r") as file:
                    text = file.read()
                    self.textEdit.clear()
                    self.textEdit.append(text)

                    self.current_file = OpenDialogData
                    self.setWindowTitle(self.app_title + f" {self.current_file}")
                    self.action_close.setEnabled(True)

        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))
    ##############################################

    def export_docx_text(self):
        try:
            SaveDialogData = QFileDialog.getSaveFileName(
                self, 'Сохранить файл', os.path.expanduser("~"), "*.docx")
            if os.name == "nt":
                SaveDialogData = SaveDialogData[0]
            elif os.name == "posix": # Fix file saving on linux
                SaveDialogData = SaveDialogData[0] + SaveDialogData[1][1::]
            # print(SaveDialogData)

            if SaveDialogData is not "" and not os.path.isfile(SaveDialogData):
                document = Document()
                text = self.textEdit.toPlainText()
                document.add_paragraph(text)
                document.save(SaveDialogData)
            elif os.path.isfile(SaveDialogData):
                self.WarningUi("Такой файл существует!")

        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))

    def import_docx_text(self):
        try:
            OpenDialogData = QFileDialog.getOpenFileName(
                self, 'Открыть файл', os.path.expanduser("~"), "*.docx")
            OpenDialogData = OpenDialogData[0]
            # print(OpenDialogData)

            if OpenDialogData is not "":
                document = Document(OpenDialogData)
                text = []

                self.current_file = OpenDialogData
                self.setWindowTitle(self.app_title + f" {self.current_file}")
                self.action_close.setEnabled(True)

                for para in document.paragraphs:
                    self.textEdit.clear()
                    self.textEdit.append(para.text)

        except Exception as err:
            # print(err)
            self.CriticalUi(str(err))

    #################################################################################################
